"""
word_generator.py
Génère un rapport Word narratif à partir des résultats d'exécution du workflow.

Utilise python-docx pour construire le document et gpt-4o-mini pour rédiger
les sections narratives à partir des sorties Python brutes.

Utilisation :
    from word_generator import generate_word_report
    doc_bytes = generate_word_report(steps, summary, template_path, file_path)
    with open("rapport.docx", "wb") as f:
        f.write(doc_bytes)
"""

from __future__ import annotations

import io
import json
import os
import re
import unicodedata
from datetime import date
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

import config

load_dotenv()

_MAX_OUTPUT_FOR_LLM = 2000  # caractères max envoyés au LLM par section


# ─────────────────────────────────────────────────────────────────────────────
# Client OpenAI
# ─────────────────────────────────────────────────────────────────────────────

def _get_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY manquante dans .env")
    return OpenAI(api_key=api_key)


def _normalize_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text or "")
    ascii_only = "".join(char for char in normalized if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]+", " ", ascii_only.lower()).strip()


def _load_blueprint(blueprint_path: Optional[str]) -> dict | None:
    if not blueprint_path:
        return None
    path = Path(blueprint_path)
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def _infer_step_kind(step_label: str) -> str | None:
    label = _normalize_text(step_label)
    if "chargement" in label:
        return "contexte"
    if "controle qualite" in label or "qualite" in label:
        return "qualite_donnees"
    if "exposition" in label:
        return "analyse_descriptive"
    if "lissage" in label:
        return "transformation"
    if "visualisation" in label:
        return "visualisation"
    if "smr" in label or "comparaison" in label:
        return "comparaison"
    if "modele" in label or "performance" in label:
        return "performance_modele"
    if "export" in label or "conclusion" in label:
        return "conclusion"
    return None


def _find_blueprint_section(step_label: str, blueprint: dict | None) -> dict | None:
    if not blueprint:
        return None

    sections = blueprint.get("sections", [])
    if not sections:
        return None

    step_kind = _infer_step_kind(step_label)
    if step_kind:
        for section in sections:
            if section.get("section_kind") == step_kind:
                return section

    normalized_label = set(_normalize_text(step_label).split())
    if not normalized_label:
        return None

    best_section = None
    best_score = 0
    for section in sections:
        title_tokens = set(_normalize_text(section.get("title", "")).split())
        score = len(normalized_label & title_tokens)
        if score > best_score:
            best_score = score
            best_section = section
    return best_section


def _blueprint_guidance_block(section: dict | None) -> str:
    if not section:
        return ""

    narrative = section.get("narrative_guidelines", [])
    tables = section.get("recommended_tables", [])
    charts = [item.get("chart_type", "") for item in section.get("chart_specs", []) if item.get("chart_type")]
    required_inputs = section.get("required_inputs", [])

    return (
        "Guidelines de rédaction validées offline :\n"
        f"- Titre de référence : {section.get('title', '')}\n"
        f"- Type de section : {section.get('section_kind', '')}\n"
        f"- Objectif : {section.get('purpose', '')}\n"
        f"- Résumé attendu : {section.get('section_description', '')}\n"
        f"- Entrées utiles : {', '.join(required_inputs) if required_inputs else 'non précisées'}\n"
        f"- Tableaux conseillés : {', '.join(tables) if tables else 'non précisés'}\n"
        f"- Graphiques conseillés : {', '.join(charts) if charts else 'non précisés'}\n"
        f"- Guidelines narratives : {' | '.join(narrative) if narrative else 'non précisées'}\n\n"
    )


def _introduction_lines(blueprint: dict | None) -> list[str]:
    default_lines = [
        "Le présent rapport présente les résultats de l'analyse de mortalité "
        "d'expérience réalisée sur le portefeuille de contrats décrit ci-dessus. "
        "L'analyse suit la méthodologie standard en sept étapes : chargement des données, "
        "contrôle qualité, calcul des expositions et taux bruts, lissage "
        "Whittaker-Henderson, visualisation, calcul du SMR et export des résultats.",
    ]
    if not blueprint:
        return default_lines

    section_titles = [section.get("title", "") for section in blueprint.get("sections", []) if section.get("title")]
    if not section_titles:
        return default_lines

    return [
        "Le présent rapport présente les résultats de l'analyse de mortalité "
        "d'expérience réalisée sur le portefeuille de contrats décrit ci-dessus.",
        "La rédaction suit un blueprint de rapport validé offline afin de conserver "
        "une structure cohérente, des messages analytiques homogènes et des éléments "
        "de preuve comparables entre études.",
        "Les sections de référence mobilisées pour ce rapport sont : "
        + ", ".join(section_titles[:8]) + ".",
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Génération de narrative par section via LLM
# ─────────────────────────────────────────────────────────────────────────────

def _narrative_for_step(step_label: str, raw_output: str, blueprint_section: dict | None = None) -> str:
    """Génère 1-3 paragraphes narratifs en français depuis la sortie brute.

    Retourne raw_output tronqué en cas d'erreur API.
    """
    truncated = raw_output[:_MAX_OUTPUT_FOR_LLM]
    if not truncated.strip():
        return "(Aucune sortie produite pour cette étape.)"

    try:
        client = _get_client()
        guidance_block = _blueprint_guidance_block(blueprint_section)
        response = client.chat.completions.create(
            model=config.FORMATTER_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Tu es un rédacteur de rapports actuariels. "
                        "Tu convertis des sorties techniques Python en prose claire "
                        "pour un lecteur non-technicien (actuaire ou direction). "
                        "Pas de code, pas de bullets, uniquement des paragraphes fluides en français."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Rédige 1 à 3 paragraphes pour la section « {step_label} » "
                        f"d'un rapport d'analyse de mortalité d'expérience.\n\n"
                        f"{guidance_block}"
                        f"Sortie technique :\n{truncated}\n\n"
                        "Présente les résultats de façon narrative, interprète les valeurs clés, "
                        "et signale toute anomalie éventuelle. "
                        "Respecte les guidelines offline quand elles sont compatibles avec les résultats réellement observés."
                    ),
                },
            ],
            max_tokens=500,
            temperature=0.3,
        )
        return (response.choices[0].message.content or "").strip()
    except Exception as exc:
        return f"{truncated}\n\n[Note : génération narrative indisponible — {exc}]"


def _narrative_summary(summary: str, blueprint: dict | None = None) -> str:
    """Reformule la synthèse en prose formelle si elle n'est pas déjà bien rédigée."""
    if not summary or not summary.strip():
        return "(Synthèse non disponible.)"
    if len(summary) > 200:
        # Déjà assez long → on l'utilise directement
        return summary
    try:
        client = _get_client()
        conclusion_section = _find_blueprint_section("conclusion", blueprint)
        guidance_block = _blueprint_guidance_block(conclusion_section)
        response = client.chat.completions.create(
            model=config.FORMATTER_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "Tu es un rédacteur de conclusions actuarielles concises et formelles en français.",
                },
                {
                    "role": "user",
                    "content": (
                        "Développe cette synthèse en 2-4 paragraphes formels pour la conclusion "
                        f"d'un rapport d'analyse de mortalité :\n\n"
                        f"{guidance_block}"
                        f"Synthèse brute :\n{summary}"
                    ),
                },
            ],
            max_tokens=400,
            temperature=0.3,
        )
        return (response.choices[0].message.content or summary).strip()
    except Exception:
        return summary


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de mise en forme Word
# ─────────────────────────────────────────────────────────────────────────────

def _set_title_style(doc: Document, file_path: str) -> None:
    """Insère la page de titre."""
    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Rapport d'analyse de mortalité d'expérience")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # bleu foncé

    # Sous-titre : nom du fichier
    if file_path:
        import os as _os
        basename = _os.path.basename(file_path)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(f"Portefeuille : {basename}")
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date de génération : {date.today().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(11)
    date_run.italic = True

    doc.add_paragraph()  # ligne vide


def _clear_document_body(doc: Document) -> None:
    """Vide le corps du document (pour un template) en gardant les styles."""
    from docx.oxml.ns import qn
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sdt"):
            body.remove(child)


def _add_section_heading(doc: Document, label: str, level: int = 1) -> None:
    doc.add_heading(label, level=level)


def _add_figure(doc: Document, fig_bytes: bytes) -> None:
    """Insère une figure PNG depuis des bytes."""
    buf = io.BytesIO(fig_bytes)
    try:
        doc.add_picture(buf, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        doc.add_paragraph("[Figure non disponible]")


# ─────────────────────────────────────────────────────────────────────────────
# Fonction principale publique
# ─────────────────────────────────────────────────────────────────────────────

def generate_word_report(
    steps: list,
    summary: str,
    template_path: Optional[str] = None,
    file_path: str = "",
    blueprint_path: Optional[str] = None,
) -> bytes:
    """Génère un rapport Word (.docx) à partir des résultats du workflow.

    Args:
        steps:         Liste de dicts {"label": str, "output": str, "figures": list[bytes]}.
        summary:       Synthèse globale (texte de l'agent ou chaîne vide).
        template_path: Chemin vers un .docx template optionnel.
                       Si fourni, le contenu existant est remplacé ; les styles sont conservés.
        file_path:     Chemin du CSV analysé (utilisé dans la page de titre).
        blueprint_path: Chemin vers un blueprint JSON validé offline. Optionnel.

    Returns:
        Contenu .docx en bytes (prêt pour dcc.send_bytes ou open("wb")).
    """
    # ── Chargement / initialisation ───────────────────────────────────────────
    if template_path:
        try:
            doc = Document(template_path)
            _clear_document_body(doc)
        except Exception:
            doc = Document()
    else:
        doc = Document()
    blueprint = _load_blueprint(blueprint_path)

    # ── Page de titre ─────────────────────────────────────────────────────────
    _set_title_style(doc, file_path)
    doc.add_page_break()

    # ── Introduction ──────────────────────────────────────────────────────────
    _add_section_heading(doc, "Introduction", level=1)
    intro_lines = _introduction_lines(blueprint)
    for line in intro_lines:
        doc.add_paragraph(line)

    # ── Sections par étape ────────────────────────────────────────────────────
    for i, step in enumerate(steps, start=1):
        label = step.get("label", f"Étape {i}")
        raw_output = step.get("output", "")
        figures = step.get("figures", [])
        skipped = step.get("status", "") == "skipped"
        blueprint_section = _find_blueprint_section(label, blueprint)

        _add_section_heading(doc, f"{label}", level=2)

        if skipped:
            doc.add_paragraph("Cette étape a été ignorée (condition métier non satisfaite).")
        else:
            narrative = _narrative_for_step(label, raw_output, blueprint_section=blueprint_section)
            for para_text in narrative.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

            for fig_bytes in figures:
                _add_figure(doc, fig_bytes)

    # ── Synthèse et conclusion ────────────────────────────────────────────────
    _add_section_heading(doc, "Synthèse et conclusion", level=1)
    narrative_summary = _narrative_summary(summary, blueprint=blueprint)
    for para_text in narrative_summary.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # ── Sérialisation en mémoire ──────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
