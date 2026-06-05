"""
_extract.py — Extraction blocs textuels depuis PDF (PyMuPDF) et DOCX (python-docx).

Sortie commune : `list[Block]` où Block = dict avec
  text          : str   — contenu paragraphe
  page          : int|None — numéro de page (PDF uniquement)
  heading_level : int|None — 1, 2, 3... ou None si paragraphe normal
  block_index   : int   — index global dans le document (0-based)

Détection heading :
  - DOCX : style "Heading N" → heading_level=N
  - PDF  : heuristique font-size (>= max_size * 0.9 + texte court < 80 chars
           ET ratio de majuscules > 0.5) OU pattern numéroté "1.2 TITRE"
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import TypedDict

log = logging.getLogger(__name__)


class Block(TypedDict, total=False):
    text: str
    page: int | None
    heading_level: int | None
    block_index: int


# Patterns nettoyage PDF
_FOOTER_RE = re.compile(r"^(page \d+|\d+\s*/\s*\d+|©.*|confidentiel.*)$", re.IGNORECASE)
_PAGE_NUM_RE = re.compile(r"^\s*\d+\s*$")
# Heading numéroté "1. TITRE" ou "1.2 Titre"
_HEADING_NUM_RE = re.compile(r"^\d+(\.\d+){0,3}\s+\S+")


def _is_likely_heading(text: str, font_size: float, max_font_size: float) -> int | None:
    """Retourne le niveau de heading (1, 2, 3) ou None. Heuristique conservative."""
    text = text.strip()
    if not text or len(text) > 120:
        return None
    # Pattern numéroté "1.2.3 Titre" → niveau = nb de segments
    m = _HEADING_NUM_RE.match(text)
    if m:
        prefix = m.group(0).split()[0]
        level = prefix.count(".") + 1
        return min(level, 3)
    # Heuristique font-size (PDF) : ≥ 90% du max + court + majuscules dominantes
    if max_font_size > 0 and font_size >= max_font_size * 0.9 and len(text) <= 80:
        upper_ratio = sum(1 for c in text if c.isupper()) / max(1, sum(1 for c in text if c.isalpha()))
        if upper_ratio >= 0.5:
            return 1
    return None


def _extract_pdf(path: Path) -> list[Block]:
    """Extrait les blocs textuels d'un PDF via PyMuPDF (fitz).

    Stratégie : page.get_text('blocks') retourne (x0, y0, x1, y1, text, block_no, type).
    On filtre les footers/numéros, détecte les headings via la taille de police
    moyenne du bloc (depuis page.get_text('dict')).
    """
    import fitz

    blocks: list[Block] = []
    bi = 0
    doc = fitz.open(str(path))
    try:
        # Pré-scan : détermine la taille de police max sur le doc (pour ratio heading)
        max_font_size = 0.0
        for page in doc:
            for b in page.get_text("dict")["blocks"]:
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        sz = span.get("size", 0.0)
                        if sz > max_font_size:
                            max_font_size = sz

        for page_no, page in enumerate(doc, start=1):
            for b in page.get_text("dict")["blocks"]:
                if b.get("type") != 0:  # 0 = texte, 1 = image
                    continue
                # Reconstitue le texte du bloc + taille moyenne
                lines_text: list[str] = []
                sizes: list[float] = []
                for line in b.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = "".join(s.get("text", "") for s in spans).strip()
                    if line_text:
                        lines_text.append(line_text)
                        sizes.extend(s.get("size", 0.0) for s in spans)
                if not lines_text:
                    continue
                text = " ".join(lines_text).strip()
                # Filtres footers et numéros de pages
                if _FOOTER_RE.match(text) or _PAGE_NUM_RE.match(text):
                    continue
                if len(text) < 3:
                    continue
                avg_size = sum(sizes) / len(sizes) if sizes else 0.0
                hlevel = _is_likely_heading(text, avg_size, max_font_size)
                blocks.append(Block(
                    text=text,
                    page=page_no,
                    heading_level=hlevel,
                    block_index=bi,
                ))
                bi += 1
    finally:
        doc.close()

    log.info("PDF %s : %d blocs extraits", path.name, len(blocks))
    return blocks


def _extract_docx(path: Path) -> list[Block]:
    """Extrait les paragraphes d'un DOCX via python-docx.

    Headings détectés via le style "Heading N" (1..9). Les tables sont aplaties
    en blocs séparés (un bloc par cellule non vide, concaténé par ligne).
    """
    from docx import Document

    doc = Document(str(path))
    blocks: list[Block] = []
    bi = 0

    # Paragraphes (ordre du document)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        hlevel: int | None = None
        if style.startswith("Heading "):
            try:
                hlevel = int(style.split(" ", 1)[1])
            except (ValueError, IndexError):
                hlevel = 1
        else:
            # Pattern numéroté fallback même sans style Heading
            m = _HEADING_NUM_RE.match(text)
            if m and len(text) <= 120:
                prefix = m.group(0).split()[0]
                hlevel = min(prefix.count(".") + 1, 3)
        blocks.append(Block(
            text=text,
            page=None,
            heading_level=hlevel,
            block_index=bi,
        ))
        bi += 1

    # Tables — aplaties en lignes de texte
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(Block(
                    text=" | ".join(cells),
                    page=None,
                    heading_level=None,
                    block_index=bi,
                ))
                bi += 1

    log.info("DOCX %s : %d blocs extraits", path.name, len(blocks))
    return blocks


def extract(path: Path) -> list[Block]:
    """Dispatch par extension. Lève ValueError si extension non supportée."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise ValueError(f"Extension non supportée : {ext} (attendu .pdf ou .docx)")
